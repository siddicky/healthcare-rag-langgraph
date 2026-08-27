import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC=Path(sys.argv[1]); OUT=Path(sys.argv[2])
RUST=RGBColor(0x63,0x13,0x00); CAMEL=RGBColor(0x9e,0x4d,0x14); CARROT_A=RGBColor(0xc2,0x39,0x0f)
BIRCH="FAF3E3"; GOLD="EDA94F"; WARM="FDF8F0"
HEAD="Quicksand"; BODY="Nunito Sans"; MONO="Roboto Mono"

doc=Document()
# page colour
bg=OxmlElement('w:background'); bg.set(qn('w:color'),WARM); doc.element.insert(0,bg)
st=doc.settings.element; d=OxmlElement('w:displayBackgroundShape'); st.append(d)

def font(run_or_style, name, size=None, bold=None, color=None):
    f=run_or_style.font; f.name=name
    rpr=f.element.rPr if hasattr(f.element,'rPr') else None
    el=run_or_style.element if hasattr(run_or_style,'element') else run_or_style._element
    rPr=el.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    for k in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'): rf.set(qn(k),name)
    if size: f.size=Pt(size)
    if bold is not None: f.bold=bold
    if color is not None: f.color.rgb=color

def para_fmt(style, before=None, after=None, line=None):
    pf=style.paragraph_format
    if before is not None: pf.space_before=Pt(before)
    if after is not None: pf.space_after=Pt(after)
    if line is not None: pf.line_spacing=line

S=doc.styles
font(S['Title'],HEAD,26,True,RUST); para_fmt(S['Title'],0,12)
S['Title'].element.get_or_add_pPr()  # drop the default bottom border
for b in S['Title'].element.pPr.findall(qn('w:pBdr')): S['Title'].element.pPr.remove(b)
font(S['Subtitle'],BODY,11,False,CAMEL); para_fmt(S['Subtitle'],0,14)
font(S['Heading 1'],HEAD,20,True,RUST); para_fmt(S['Heading 1'],24,8)
font(S['Heading 2'],HEAD,15,True,RUST); para_fmt(S['Heading 2'],18,6)
font(S['Heading 3'],BODY,11,True,CAMEL); para_fmt(S['Heading 3'],12,2)
font(S['Normal'],BODY,11,False,RUST); para_fmt(S['Normal'],0,8,1.5)
for n in ('List Bullet','List Number'):
    font(S[n],BODY,11,False,RUST); para_fmt(S[n],0,4,1.5)

def shade(run,hexfill):
    rPr=run._element.get_or_add_rPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexfill); rPr.append(sh)

def add_hyperlink(par,text,url):
    part=par.part; rid=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid)
    r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    rf=OxmlElement('w:rFonts'); [rf.set(qn(k),BODY) for k in ('w:ascii','w:hAnsi','w:cs')]; rPr.append(rf)
    c=OxmlElement('w:color'); c.set(qn('w:val'),'C2390F'); rPr.append(c)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    r.append(rPr); t=OxmlElement('w:t'); t.text=text; t.set(qn('xml:space'),'preserve'); r.append(t); h.append(r); par._p.append(h)

INLINE=re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
def inline(par,text):
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith('**'):
            r=par.add_run(tok[2:-2]); r.bold=True
        elif tok.startswith('`'):
            r=par.add_run(tok[1:-1]); font(r,MONO,10,None,RUST); shade(r,BIRCH)
        elif tok.startswith('['):
            m=re.match(r'\[([^\]]+)\]\(([^)]+)\)',tok); add_hyperlink(par,m.group(1),m.group(2))
        else:
            par.add_run(tok)

lines=SRC.read_text().splitlines()
i=0; first_para_done=False
while i<len(lines):
    ln=lines[i]
    if not ln.strip(): i+=1; continue
    if ln.startswith('# '):
        doc.add_paragraph(ln[2:],style='Title'); i+=1
        # next non-empty paragraph is the byline → Subtitle
        while i<len(lines) and not lines[i].strip(): i+=1
        if i<len(lines) and not lines[i].startswith(('#','-','*')):
            p=doc.add_paragraph(style='Subtitle'); inline(p,lines[i]); i+=1
        continue
    if ln.startswith('## '): doc.add_paragraph(ln[3:],style='Heading 1'); i+=1; continue
    if ln.startswith('### '): doc.add_paragraph(ln[4:],style='Heading 2'); i+=1; continue
    if ln.strip()=='---': i+=1; continue
    m=re.match(r'^\*\*([^*]+)\*\*\s*(.*)$',ln)
    if m and (not m.group(2) or m.group(1).endswith('.')):
        doc.add_paragraph(m.group(1).rstrip('.'),style='Heading 3')
        if m.group(2): p=doc.add_paragraph(); inline(p,m.group(2))
        i+=1; continue
    if re.match(r'^- ',ln): p=doc.add_paragraph(style='List Bullet'); inline(p,ln[2:]); i+=1; continue
    if re.match(r'^\d+\. ',ln): p=doc.add_paragraph(style='List Number'); inline(p,re.sub(r'^\d+\. ','',ln)); i+=1; continue
    p=doc.add_paragraph(); inline(p,ln); i+=1

# header wordmark
hp=doc.sections[0].header.paragraphs[0]; r=hp.add_run('nymble health'); font(r,HEAD,10,True,RUST)
r2=hp.add_run('   healthcare-rag · technical write-up'); font(r2,BODY,9,False,CAMEL)
doc.save(OUT); print('wrote',OUT, OUT.stat().st_size//1024,'KB')
